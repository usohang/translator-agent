"""IR → DOCX 생성. python-docx 기반, 8장 스타일 규칙 적용."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.ir import IRDocument, Block, BlockType
from ..pagemap.mapper import build_page_map

# ── 스타일 상수 (기획서 8.1) ──────────────────────────────
BODY_PT = 11
H1_PT, H2_PT, H3_PT = 22, 17, 13
CAPTION_PT = 9
PAGE_MARK_PT = 9
GRAY = RGBColor(0x88, 0x88, 0x88)
FONT_NAME = "맑은 고딕"


def build_docx(ir: IRDocument, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_page_map(ir)

    doc = Document()
    _set_default_style(doc)

    prev_src_page = 0

    for block in ir.blocks:
        text = block.text_tgt or block.text_src

        # 원문 페이지 표기 (페이지가 바뀌면 삽입)
        if block.original_page and block.original_page != prev_src_page:
            _add_page_marker(doc, block.original_page)
            prev_src_page = block.original_page

        if block.type == BlockType.HEADING:
            _add_heading(doc, text, block.level or 1)
        elif block.type == BlockType.PARAGRAPH:
            _add_paragraph(doc, text)
        elif block.type == BlockType.TABLE and block.table:
            _add_table(doc, block)
        elif block.type == BlockType.IMAGE and block.image_ref:
            _add_image(doc, block.image_ref)
        elif block.type == BlockType.CAPTION:
            _add_caption(doc, text)
        elif block.type == BlockType.LIST:
            _add_list_item(doc, text)

    # 문서 끝 페이지 매핑 표
    if ir.page_map:
        _add_page_map_table(doc, ir)

    doc.save(str(output_path))
    return output_path


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(BODY_PT)


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: H1_PT, 2: H2_PT, 3: H3_PT}
    pt = sizes.get(level, H2_PT)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(pt)
    run.font.name = FONT_NAME


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(BODY_PT)
    run.font.name = FONT_NAME


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(CAPTION_PT)
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME


def _add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(BODY_PT)
    run.font.name = FONT_NAME


def _add_page_marker(doc: Document, page_num: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"〔원문 p.{page_num}〕")
    run.font.size = Pt(PAGE_MARK_PT)
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME


def _add_image(doc: Document, image_ref: str) -> None:
    path = Path(image_ref)
    if not path.exists():
        return
    try:
        doc.add_picture(str(path), width=Inches(5.5))
    except Exception:
        _add_caption(doc, f"[이미지: {path.name}]")


def _add_table(doc: Document, block: Block) -> None:
    tdata = block.table
    if not tdata:
        return

    all_rows = ([tdata.headers] if tdata.headers else []) + tdata.rows
    if not all_rows:
        return

    ncols = max(len(r) for r in all_rows)
    table = doc.add_table(rows=len(all_rows), cols=ncols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(all_rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(BODY_PT)
            run.font.name = FONT_NAME
            if r_idx == 0:
                run.bold = True


def _add_page_map_table(doc: Document, ir: IRDocument) -> None:
    doc.add_page_break()
    _add_heading(doc, "원문 ↔ 번역본 페이지 매핑표", level=2)

    headers = ["번역본 페이지", "원문 페이지", "챕터/섹션"]
    rows = [[e.tgt_page, e.src_page, e.section] for e in ir.page_map]

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    for c_idx, h in enumerate(headers):
        cell = table.cell(0, c_idx)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(BODY_PT)
        run.font.name = FONT_NAME

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(BODY_PT)
            run.font.name = FONT_NAME
